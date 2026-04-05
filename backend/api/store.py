from __future__ import annotations

import asyncio
import base64
import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Literal
from uuid import UUID

import asyncpg

from .schemas import DocumentContentResponse, DocumentExportResult, DocumentVersionItem

DEFAULT_DOCUMENT_CONTENT = (
    "The quick brown fox jumps over the lazy dog. "
    "This is a sample document for testing the editor.\n\n"
    'Try selecting some text and clicking "Rewrite" to see the AI assistant in action.\n\n'
    "You can edit this content freely, and the version control system will track changes."
)

ExportFormat = Literal["md", "pdf", "docx"]


@dataclass(slots=True)
class VersionEntry:
    version_id: int
    content: str
    created_at: datetime
    created_by: str


@dataclass(slots=True)
class DocumentState:
    id: str
    title: str
    content: str
    version_id: int
    last_modified: datetime
    versions: list[VersionEntry] = field(default_factory=list)

    def as_content_response(self) -> DocumentContentResponse:
        return DocumentContentResponse(
            id=self.id,
            title=self.title,
            content=self.content,
            versionId=self.version_id,
            lastModified=self.last_modified.isoformat(),
        )


class VersionConflictError(Exception):
    pass


class InMemoryDocumentStore:
    def __init__(self) -> None:
        self._lock = asyncio.Lock()
        self._documents: dict[str, DocumentState] = {}
        self._pool: asyncpg.Pool | None = None

    async def bind_pool(self, pool: asyncpg.Pool | None) -> None:
        self._pool = pool
        if pool is None:
            return
        async with pool.acquire() as conn:
            await conn.execute(
                """
                CREATE TABLE IF NOT EXISTS document_live_content (
                    doc_id UUID PRIMARY KEY REFERENCES documents (id) ON DELETE CASCADE,
                    content TEXT NOT NULL DEFAULT '',
                    version_id INT NOT NULL DEFAULT 1,
                    updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """
            )
            await conn.execute(
                """
                CREATE TABLE IF NOT EXISTS document_text_versions (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    doc_id UUID NOT NULL REFERENCES documents (id) ON DELETE CASCADE,
                    version_number INT NOT NULL,
                    content TEXT NOT NULL,
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                    created_by UUID NOT NULL REFERENCES users (id) ON DELETE RESTRICT,
                    CONSTRAINT uq_document_text_versions UNIQUE (doc_id, version_number)
                )
                """
            )

    async def ensure_document(
        self,
        doc_id: str,
        title: str,
        created_by: str,
        initial_content: str = DEFAULT_DOCUMENT_CONTENT,
    ) -> None:
        async with self._lock:
            if self._pool is None:
                if doc_id not in self._documents:
                    now = datetime.now(timezone.utc)
                    self._documents[doc_id] = DocumentState(
                        id=doc_id,
                        title=title,
                        content=initial_content,
                        version_id=1,
                        last_modified=now,
                        versions=[
                            VersionEntry(
                                version_id=1,
                                content=initial_content,
                                created_at=now,
                                created_by=created_by,
                            )
                        ],
                    )
                else:
                    self._documents[doc_id].title = title
                return

        assert self._pool is not None
        async with self._pool.acquire() as conn:
            await conn.execute(
                """
                INSERT INTO document_live_content (doc_id, content, version_id, updated_at)
                VALUES ($1, $2, 1, NOW())
                ON CONFLICT (doc_id) DO NOTHING
                """,
                UUID(doc_id),
                initial_content,
            )
            await conn.execute(
                """
                INSERT INTO document_text_versions (doc_id, version_number, content, created_by)
                SELECT $1, 1, $2, $3
                WHERE NOT EXISTS (
                    SELECT 1 FROM document_text_versions WHERE doc_id = $1 AND version_number = 1
                )
                """,
                UUID(doc_id),
                initial_content,
                UUID(created_by),
            )

    async def get_document(self, doc_id: str) -> DocumentContentResponse:
        if self._pool is None:
            async with self._lock:
                document = self._documents.get(doc_id)
                if document is None:
                    raise KeyError(doc_id)
                return document.as_content_response()

        assert self._pool is not None
        async with self._pool.acquire() as conn:
            row = await conn.fetchrow(
                """
                SELECT d.id, d.title, c.content, c.version_id, c.updated_at
                FROM documents d
                JOIN document_live_content c ON c.doc_id = d.id
                WHERE d.id = $1 AND d.is_deleted = FALSE
                """,
                UUID(doc_id),
            )
        if row is None:
            raise KeyError(doc_id)
        return DocumentContentResponse(
            id=str(row["id"]),
            title=row["title"],
            content=row["content"],
            versionId=int(row["version_id"]),
            lastModified=row["updated_at"].isoformat(),
        )

    async def get_version(self, doc_id: str) -> int:
        return (await self.get_document(doc_id)).versionId

    async def update_document(
        self,
        doc_id: str,
        content: str,
        version_id: int,
        user_id: str,
    ) -> DocumentContentResponse:
        if self._pool is None:
            async with self._lock:
                document = self._documents.get(doc_id)
                if document is None:
                    raise KeyError(doc_id)
                if version_id != document.version_id:
                    raise VersionConflictError(
                        f"Version conflict: expected {document.version_id}, got {version_id}"
                    )

                document.version_id += 1
                document.content = content
                document.last_modified = datetime.now(timezone.utc)
                document.versions.append(
                    VersionEntry(
                        version_id=document.version_id,
                        content=content,
                        created_at=document.last_modified,
                        created_by=user_id,
                    )
                )
                return document.as_content_response()

        assert self._pool is not None
        async with self._pool.acquire() as conn:
            async with conn.transaction():
                live_row = await conn.fetchrow(
                    """
                    SELECT version_id
                    FROM document_live_content
                    WHERE doc_id = $1
                    FOR UPDATE
                    """,
                    UUID(doc_id),
                )
                if live_row is None:
                    raise KeyError(doc_id)
                current_version = int(live_row["version_id"])
                if current_version != version_id:
                    raise VersionConflictError(
                        f"Version conflict: expected {current_version}, got {version_id}"
                    )

                next_version = current_version + 1
                await conn.execute(
                    """
                    UPDATE document_live_content
                    SET content = $2, version_id = $3, updated_at = NOW()
                    WHERE doc_id = $1
                    """,
                    UUID(doc_id),
                    content,
                    next_version,
                )
                await conn.execute(
                    """
                    INSERT INTO document_text_versions (doc_id, version_number, content, created_by)
                    VALUES ($1, $2, $3, $4)
                    """,
                    UUID(doc_id),
                    next_version,
                    content,
                    UUID(user_id),
                )

            row = await conn.fetchrow(
                """
                SELECT d.id, d.title, c.content, c.version_id, c.updated_at
                FROM documents d
                JOIN document_live_content c ON c.doc_id = d.id
                WHERE d.id = $1
                """,
                UUID(doc_id),
            )

        return DocumentContentResponse(
            id=str(row["id"]),
            title=row["title"],
            content=row["content"],
            versionId=int(row["version_id"]),
            lastModified=row["updated_at"].isoformat(),
        )

    async def list_versions(self, doc_id: str) -> list[DocumentVersionItem]:
        if self._pool is None:
            async with self._lock:
                document = self._documents.get(doc_id)
                if document is None:
                    raise KeyError(doc_id)
                return [
                    DocumentVersionItem(
                        version_id=entry.version_id,
                        created_at=entry.created_at.isoformat(),
                        created_by=entry.created_by,
                    )
                    for entry in sorted(document.versions, key=lambda item: item.version_id, reverse=True)
                ]

        assert self._pool is not None
        async with self._pool.acquire() as conn:
            rows = await conn.fetch(
                """
                SELECT version_number, created_at, created_by
                FROM document_text_versions
                WHERE doc_id = $1
                ORDER BY version_number DESC
                """,
                UUID(doc_id),
            )
        return [
            DocumentVersionItem(
                version_id=int(row["version_number"]),
                created_at=row["created_at"].isoformat(),
                created_by=str(row["created_by"]),
            )
            for row in rows
        ]

    async def create_snapshot(self, doc_id: str, user_id: str, snapshot: str | None = None) -> DocumentVersionItem:
        current = await self.get_document(doc_id)
        content = current.content
        if snapshot:
            try:
                content = base64.b64decode(snapshot.encode("utf-8")).decode("utf-8")
            except Exception:
                content = current.content
        updated = await self.update_document(doc_id, content, current.versionId, user_id)
        return DocumentVersionItem(
            version_id=updated.versionId,
            created_at=updated.lastModified,
            created_by=user_id,
        )

    async def revert_document(self, doc_id: str, version_id: int, user_id: str) -> DocumentVersionItem:
        if self._pool is None:
            async with self._lock:
                document = self._documents.get(doc_id)
                if document is None:
                    raise KeyError(doc_id)
                target = next((entry for entry in document.versions if entry.version_id == version_id), None)
                if target is None:
                    raise KeyError(version_id)
                updated = await self.update_document(doc_id, target.content, document.version_id, user_id)
                return DocumentVersionItem(
                    version_id=updated.versionId,
                    created_at=updated.lastModified,
                    created_by=user_id,
                )

        assert self._pool is not None
        async with self._pool.acquire() as conn:
            row = await conn.fetchrow(
                """
                SELECT content
                FROM document_text_versions
                WHERE doc_id = $1 AND version_number = $2
                """,
                UUID(doc_id),
                version_id,
            )
        if row is None:
            raise KeyError(version_id)
        current = await self.get_document(doc_id)
        updated = await self.update_document(doc_id, row["content"], current.versionId, user_id)
        return DocumentVersionItem(
            version_id=updated.versionId,
            created_at=updated.lastModified,
            created_by=user_id,
        )

    async def export_document(self, doc_id: str, fmt: ExportFormat) -> DocumentExportResult:
        document = await self.get_document(doc_id)

        if fmt == "md":
            markdown = f"# {document.title or 'Untitled Document'}\n\n{document.content}\n"
            return DocumentExportResult(
                filename=f"{document.id}.md",
                media_type="text/markdown; charset=utf-8",
                content=markdown.encode("utf-8"),
            )

        if fmt == "docx":
            from docx import Document as DocxDocument

            buffer = io.BytesIO()
            docx = DocxDocument()
            docx.add_heading(document.title or "Untitled Document", level=1)
            for block in document.content.split("\n\n"):
                docx.add_paragraph(block)
            docx.save(buffer)
            return DocumentExportResult(
                filename=f"{document.id}.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=buffer.getvalue(),
            )

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        from reportlab.pdfgen import canvas

        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - inch
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(inch, y, document.title or "Untitled Document")
        y -= 0.5 * inch
        pdf.setFont("Helvetica", 11)
        for paragraph in document.content.split("\n\n"):
            for line in paragraph.splitlines() or [""]:
                if y <= inch:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 11)
                    y = height - inch
                pdf.drawString(inch, y, line)
                y -= 14
            y -= 10
        pdf.save()
        return DocumentExportResult(
            filename=f"{document.id}.pdf",
            media_type="application/pdf",
            content=buffer.getvalue(),
        )
